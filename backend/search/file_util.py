import fitz  # PyMuPDF
import docx  # python-docx
import pandas as pd
from io import BytesIO
from PIL import Image
import base64

def process_uploaded_file(file_storage):
    file_name = file_storage.filename
    file_type = file_storage.content_type or ''
    
    # Get file extension
    file_extension = file_name.lower().split('.')[-1] if '.' in file_name else ''
    
    try:
        # Read file content as bytes
        content_bytes = file_storage.read()
        
        # Process based on file extension
        if file_extension == 'pdf':
            return process_pdf(content_bytes)
        elif file_extension == 'docx':
            return process_docx(content_bytes)
        elif file_extension in ['xlsx', 'xls']:
            return process_excel(content_bytes)
        elif file_type.startswith('image/'):
            image_result = process_image(file_storage)
            if isinstance(image_result, dict):
                return f"[Image file: {file_name} - Image content available for AI analysis]"
            else:
                return image_result  # Error message
        else:
            # Try to read as plain text
            return process_text(content_bytes, file_name)
            
    except Exception as e:
        return f"Error processing file {file_name}: {str(e)}"

def process_pdf(content_bytes):
    """Extract text from PDF file using PyMuPDF."""
    try:
        # Open PDF from bytes
        pdf_document = fitz.open(stream=content_bytes, filetype="pdf")
        
        text_content = ""
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text_content += f"\n--- Page {page_num + 1} ---\n"
            text_content += page.get_text()
        
        pdf_document.close()
        return text_content.strip()
    except Exception as e:
        raise Exception(f"Failed to process PDF: {str(e)}")

def process_docx(content_bytes):
    """Extract text from DOCX file."""
    try:
        docx_file = BytesIO(content_bytes)
        doc = docx.Document(docx_file)
        
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            text_content += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text_content += " | ".join(row_text) + "\n"
        
        return text_content.strip()
    except Exception as e:
        raise Exception(f"Failed to process DOCX: {str(e)}")

def process_excel(content_bytes):
    """Extract data from Excel file."""
    try:
        excel_file = BytesIO(content_bytes)
        
        # Read all sheets
        excel_data = pd.read_excel(excel_file, sheet_name=None)
        
        text_content = ""
        for sheet_name, df in excel_data.items():
            text_content += f"\n--- Sheet: {sheet_name} ---\n"
            
            # Convert DataFrame to string representation
            text_content += df.to_string(index=False, na_rep='')
            text_content += "\n"
        
        return text_content.strip()
    except Exception as e:
        raise Exception(f"Failed to process Excel: {str(e)}")

def process_text(content_bytes, file_name):
    """Process as plain text file."""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text_content = content_bytes.decode(encoding)
                return text_content
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, return error
        raise Exception("Unable to decode file with common text encodings")
        
    except Exception as e:
        raise Exception(f"Failed to process text file: {str(e)}")

def process_image(image_file):
    try:
        # Reset file pointer to beginning
        image_file.seek(0)
        
        # Read image bytes
        img_bytes = image_file.read()
        
        # Check if file is empty
        if not img_bytes:
            return f"[Error: Image file {getattr(image_file, 'filename', 'unknown')} is empty]"
        
        # Reset file pointer again for PIL
        image_file.seek(0)
        
        # Try to open and validate the image
        try:
            with Image.open(image_file) as img:
                # Verify it's a valid image
                img.verify()
                
                # Get format and mime type
                img_format = img.format
                if img_format:
                    mime_type = f"image/{img_format.lower()}"
                else:
                    # Fallback mime type detection
                    mime_type = "image/jpeg"  # Default fallback
                
        except Exception as img_error:
            return f"[Error: Cannot process image file {getattr(image_file, 'filename', 'unknown')}: {str(img_error)}]"
        
        # Reset and read again for base64 encoding (verify() closes the image)
        image_file.seek(0)
        img_bytes = image_file.read()
        
        # Encode to base64
        img_base64 = base64.b64encode(img_bytes).decode("utf-8")
        
        # Create content block
        content_block = {
            "type": "image_url",
            "image_url": {"url": f"data:{mime_type};base64,{img_base64}"}
        }
        
        return content_block
        
    except Exception as e:
        return f"[Error processing image file {getattr(image_file, 'filename', 'unknown')}: {str(e)}]"
